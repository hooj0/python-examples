#!/usr/bin/env python3
# encoding: utf-8
# @author: hoojo
# @email: hoojo_@126.com
# @github: https://github.com/hooj0
# @create date: 2020-03-04
# @copyright by hoojo @2020
# @changelog Generator SQL shell to Office word document
import os
import re
from docx import Document

# ===============================================================================
#     Generator SQL shell to Office word document
# ===============================================================================
# 描述：读取sql脚本，生成word文档
# -------------------------------------------------------------------------------


# -------------------------------------------------------------------------------
# 读取SQL脚本生成 word的文档 工具类
# -------------------------------------------------------------------------------
class SqlFile2WordUtils:

    def read_sql_file(self, sql_file=None):

        if sql_file is None:
            print("sql file path 为空")
        elif not os.path.exists(sql_file):
            print('%s is not found!' % sql_file)
            raise IOError('%s is not found!' % sql_file)

        records = {}
        with open(sql_file, 'r', encoding=u'utf-8') as file:
            line = file.readline()

            table_name, keys = None, []
            while len(line) > 0:
                line = line.strip()
                match_tab_name = re.match(r'CREATE\s+TABLE\s+(.*)', line, re.M | re.I)

                if match_tab_name:
                    table_name = re.sub(r'\W', "", match_tab_name.group(1)).lower()
                    records[table_name] = []
                elif table_name:
                    if line.endswith(");"):
                        records[table_name][0] = records[table_name][0][:-1] + (" ".join(keys),)
                        records[table_name] = tuple(records[table_name])
                        table_name = None
                        keys.clear()
                    else:
                        line = line.lower()
                        fields = re.split(r'\s+', line, re.M | re.I)

                        if line.find("primary key") != -1:
                            keys.append(line.replace("primary key", "pk"))
                        elif line.find("foreign key") != -1:
                            keys.append(line.replace("foreign key", "fk"))
                        elif line.find("references") != -1:
                            keys.append(line)
                        elif len(fields) > 1:
                            field_name = re.sub(r'\W', "", fields[0])
                            record = (field_name, fields[1].upper(), re.sub(r'\D', "", line), "")

                            records[table_name].append(record)

                line = file.readline()

        return records

    def add_table(self, records, doc_path):
        """
            records = {
                'tab1': ((1, 'aaa', 'bbb', "zzz"), (2, 'bbb', 'ccc', "zzz")),
                'tab2': ((1, 'aaa', 'bbb', "zzz"), (2, 'bbb', 'ccc', "zzz")),
                'tab3': ((1, 'aaa', 'bbb', "zzz"), (2, 'bbb', 'ccc', "zzz"))
            }
        """
        document = Document()

        title = document.add_paragraph()
        r_total = title.add_run(u'生成表格：')
        r_total.font.bold = True

        for name, record in records.items():
            document.add_heading(name, level=2)

            table = document.add_table(1, 4, style="Table Grid")

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = u'字段名称'
            hdr_cells[1].text = u'字段类型'
            hdr_cells[2].text = u'长度范围'
            hdr_cells[3].text = u'字段说明'

            for field, t, length, remark in record:
                cells = table.add_row().cells

                cells[0].text = str(field)
                cells[1].text = str(t)
                cells[2].text = str(length)
                cells[3].text = str(remark)

            document.add_paragraph()

        document.save(doc_path)

    def transform(self, sql_path, doc_path):
        records = self.read_sql_file(sql_path)
        self.add_table(records, doc_path)


if __name__ == "__main__":
    util = SqlFile2WordUtils()
    util.transform("quartz.sql", "sql.doc")
    #util.add_table("a.docx")
    #print(util.read_sql_file("quartz.sql"))
